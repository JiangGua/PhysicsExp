#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches

# docu = Document()

# docu.add_heading('Docu Title', 0)

# p = docu.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
# docu.add_heading('Heading, level 1', level=1)
# docu.add_paragraph('Intense quote', style='IntenseQuote')
 
# p11 = docu.add_paragraph('A plain paragraph having some ')
# p2 = docu.add_paragraph('A plain paragraph having some ')
# p3 = docu.add_paragraph('A plain paragraph having some ')

# docu.add_paragraph(
    # 'first item in unordered list', style='ListBullet'
# )
# docu.add_paragraph(
    # 'first item in ordered list', style='ListNumber'
# )
 
# # docu.add_picture('monty-truth.png', width=Inches(1.25))
 
# table = docu.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'

# docu.add_picture('./pic.png')

# docu.save('gen.docx')


def gendocx(name, *elem):
    imglist = ['.png', '.jpg', '.bmp']
    docu = Document()
    p = docu.add_paragraph('')
    for i in elem:
        if i[-4:] in imglist:
            docu.add_picture(i, width=Inches(5.))
            p = docu.add_paragraph('')
        else:
            p.add_run(i)
    docu.save(name)


